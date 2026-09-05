"""
generate_doc.py
Generates a Word (.docx) document with all project source code files.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"D:\PPE"

# All files to include, in order
FILES = [
    # ── Root files ─────────────────────────────────────────────────────────────
    ("package.json",             "Root Config",          r"package.json"),
    ("construction_safety.sql",  "Database Schema & Seed", r"construction_safety.sql"),

    # ── Backend ────────────────────────────────────────────────────────────────
    ("backend/server.js",                 "Backend – Entry Point",   r"backend\server.js"),
    ("backend/config/database.js",        "Backend – Config",        r"backend\config\database.js"),
    ("backend/config/upload.js",          "Backend – Config",        r"backend\config\upload.js"),
    ("backend/utils/cron.js",             "Backend – Utils",         r"backend\utils\cron.js"),
    ("backend/utils/payroll.js",          "Backend – Utils",         r"backend\utils\payroll.js"),
    ("backend/routes/auth.js",            "Backend – Routes",        r"backend\routes\auth.js"),
    ("backend/routes/workers.js",         "Backend – Routes",        r"backend\routes\workers.js"),
    ("backend/routes/attendance.js",      "Backend – Routes",        r"backend\routes\attendance.js"),
    ("backend/routes/violations.js",      "Backend – Routes",        r"backend\routes\violations.js"),
    ("backend/routes/salary.js",          "Backend – Routes",        r"backend\routes\salary.js"),
    ("backend/routes/health.js",          "Backend – Routes",        r"backend\routes\health.js"),
    ("backend/routes/face_attendance.js", "Backend – Routes",        r"backend\routes\face_attendance.js"),
    ("backend/routes/cameras.js",         "Backend – Routes",        r"backend\routes\cameras.js"),
    ("backend/routes/fines.js",           "Backend – Routes",        r"backend\routes\fines.js"),
    ("backend/alter_db.js",               "Backend – Utilities",     r"backend\alter_db.js"),

    # ── Face Recognition / AI ──────────────────────────────────────────────────
    ("Face_recognition/requirements.txt",   "AI Module – Dependencies",  r"Face_recognition\requirements.txt"),
    ("Face_recognition/app.py",             "AI Module – Face Auth API", r"Face_recognition\app.py"),
    ("Face_recognition/intelligent_cctv.py","AI Module – CCTV Engine",   r"Face_recognition\intelligent_cctv.py"),
    ("Face_recognition/camera_stream.py",   "AI Module – Camera Stream", r"Face_recognition\camera_stream.py"),
    ("Face_recognition/recognize.py",       "AI Module – Utilities",     r"Face_recognition\recognize.py"),
    ("Face_recognition/register.py",        "AI Module – Utilities",     r"Face_recognition\register.py"),
    ("Face_recognition/rename.py",          "AI Module – Utilities",     r"Face_recognition\rename.py"),
    ("Face_recognition/cleanup_photos.py",  "AI Module – Utilities",     r"Face_recognition\cleanup_photos.py"),

    # ── Frontend ───────────────────────────────────────────────────────────────
    ("frontend/auth.js",              "Frontend – Auth Guard",   r"frontend\auth.js"),
    ("frontend/landingpage.html",     "Frontend – Pages",        r"frontend\landingpage.html"),
    ("frontend/login.html",           "Frontend – Pages",        r"frontend\login.html"),
    ("frontend/signup.html",          "Frontend – Pages",        r"frontend\signup.html"),
    ("frontend/dashboard.html",       "Frontend – Pages",        r"frontend\dashboard.html"),
    ("frontend/workers.html",         "Frontend – Pages",        r"frontend\workers.html"),
    ("frontend/attendance.html",      "Frontend – Pages",        r"frontend\attendance.html"),
    ("frontend/violations.html",      "Frontend – Pages",        r"frontend\violations.html"),
    ("frontend/salary.html",          "Frontend – Pages",        r"frontend\salary.html"),
    ("frontend/health.html",          "Frontend – Pages",        r"frontend\health.html"),
    ("frontend/cctv.html",            "Frontend – Pages",        r"frontend\cctv.html"),
    ("frontend/face-recognition.html","Frontend – Pages",        r"frontend\face-recognition.html"),
    ("frontend/camera.html",          "Frontend – Pages",        r"frontend\camera.html"),
]


def set_cell_background(cell, hex_color):
    """Set a table-cell background color (e.g., '2B2B2B')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_code_block(doc, code_text):
    """Add a styled code block inside a table cell."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_background(cell, '1E1E1E')  # dark background

    # Clear default paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    lines = code_text.split('\n')
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)  # light grey text


def main():
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover / Title ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Construction Site Safety Management System")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Final Year Project – Complete Source Code Documentation")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    desc = doc.add_paragraph(
        "This document contains the complete, unabridged source code of every file in the "
        "Construction Site Safety Management System FYP. The project integrates a Node.js/Express "
        "REST API backend, a Python Flask AI module (YOLO + DeepFace), and an HTML/CSS/JavaScript "
        "frontend into a unified construction-safety platform."
    )
    desc.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────────────────
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    current_section = None
    section_num = 0
    file_num = 0
    for display_name, section_label, rel_path in FILES:
        if section_label != current_section:
            current_section = section_label
            section_num += 1
            p = doc.add_paragraph(style='List Number')
            p.clear()
            run = p.add_run(f"  {section_label}")
            run.bold = True
            run.font.size = Pt(11)
        file_num += 1
        fp = doc.add_paragraph(style='List Bullet')
        fp.clear()
        fr = fp.add_run(f"        {display_name}")
        fr.font.size = Pt(10)

    doc.add_page_break()

    # ── Per-file sections ──────────────────────────────────────────────────────
    current_section = None
    for display_name, section_label, rel_path in FILES:
        full_path = os.path.join(BASE_DIR, rel_path)

        # Section heading when it changes
        if section_label != current_section:
            current_section = section_label
            sh = doc.add_heading(section_label, level=1)
            sh.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        # File heading
        fh = doc.add_heading(display_name, level=2)
        fh.runs[0].font.color.rgb = RGBColor(0x0A, 0x55, 0xBB)

        # Path label
        path_p = doc.add_paragraph()
        path_r = path_p.add_run(f"Path: {display_name}")
        path_r.font.size   = Pt(9)
        path_r.font.italic = True
        path_r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        path_p.paragraph_format.space_after = Pt(4)

        # Read file
        if os.path.exists(full_path):
            try:
                with open(full_path, 'r', encoding='utf-8', errors='replace') as f:
                    code = f.read()
                add_code_block(doc, code)
            except Exception as e:
                doc.add_paragraph(f"[ERROR reading file: {e}]")
        else:
            doc.add_paragraph(f"[FILE NOT FOUND: {full_path}]")

        doc.add_paragraph()  # spacing after each file
        doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────────
    out_path = os.path.join(BASE_DIR, "FYP_Complete_Source_Code.docx")
    doc.save(out_path)
    print(f"\n[OK] Document saved to: {out_path}")
    print(f"   Total files included: {len(FILES)}")


if __name__ == '__main__':
    main()
